"""
Resume parsing service for extracting and analyzing candidate resumes.
Uses pdfplumber for PDFs, python-docx for Word documents, and Claude for intelligent parsing.
Optimized for speed and accuracy.
"""

import io
import json
import logging
import httpx
import asyncio
from typing import Optional
from ..config import settings
from ..schemas.candidate import ParsedResume, ExperienceItem, EducationItem, ProjectItem, PersonalizedQuestion

logger = logging.getLogger("pathway.resume")


class ResumeService:
    """Service for parsing resumes and generating personalized questions using Claude."""

    def __init__(self):
        self.anthropic_api_key = settings.anthropic_api_key
        self.claude_model = settings.claude_model
        self.anthropic_base_url = "https://api.anthropic.com/v1"
        # Fallback to DeepSeek if Claude not configured
        self.deepseek_api_key = settings.deepseek_api_key
        self.deepseek_base_url = settings.deepseek_base_url

    async def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract raw text from a resume file (PDF or Word).

        Args:
            file_bytes: Raw file bytes
            filename: Original filename (used to determine file type)

        Returns:
            Extracted text content
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return self._extract_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            return self._extract_from_docx(file_bytes)
        elif filename_lower.endswith('.doc'):
            raise ValueError("Old .doc format not supported. Please use .docx or PDF.")
        else:
            raise ValueError(f"Unsupported file format: {filename}. Please upload PDF or DOCX.")

    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")

        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from Word document using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    async def _call_claude(self, system_prompt: str, user_prompt: str, max_tokens: int = 3000) -> dict:
        """
        Call Claude API for fast, accurate AI processing.

        Args:
            system_prompt: System instructions
            user_prompt: User message
            max_tokens: Maximum response tokens

        Returns:
            Parsed JSON response
        """
        async with httpx.AsyncClient(timeout=30.0) as client:  # Reduced timeout for speed
            response = await client.post(
                f"{self.anthropic_base_url}/messages",
                headers={
                    "x-api-key": self.anthropic_api_key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.claude_model,
                    "max_tokens": max_tokens,
                    "system": system_prompt,
                    "messages": [
                        {"role": "user", "content": user_prompt}
                    ]
                }
            )
            response.raise_for_status()
            result = response.json()
            content = result["content"][0]["text"]

            # Extract JSON from response (Claude may include markdown code blocks)
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]

            return json.loads(content.strip())

    async def parse_resume(self, raw_text: str) -> ParsedResume:
        """
        Parse raw resume text into structured data using Claude LLM.
        Optimized for speed and accuracy.

        Args:
            raw_text: Extracted text from resume

        Returns:
            ParsedResume with structured data
        """
        if not self.anthropic_api_key:
            # Fallback to DeepSeek if Claude not configured
            return await self._parse_resume_deepseek(raw_text)

        system_prompt = """You are an expert resume parser. Extract structured information from resumes with high accuracy.
Be precise and only include information clearly stated. For missing fields, use null.
Use date formats like "2020-01", "2020", or "Present".
Respond with valid JSON only - no markdown, no explanations."""

        user_prompt = f"""Parse this resume into structured JSON:

{raw_text[:10000]}

Return this exact JSON structure:
{{
    "name": "Full Name or null",
    "email": "email@example.com or null",
    "phone": "phone number or null",
    "location": "City, State/Country or null",
    "summary": "Professional summary or null",
    "skills": ["skill1", "skill2"],
    "experience": [
        {{
            "company": "Company Name",
            "title": "Job Title",
            "start_date": "2020-01",
            "end_date": "Present",
            "description": "Brief description",
            "highlights": ["achievement 1", "achievement 2"]
        }}
    ],
    "education": [
        {{
            "institution": "University Name",
            "degree": "Degree Type",
            "field_of_study": "Major or null",
            "start_date": "2016",
            "end_date": "2020",
            "gpa": "3.8 or null"
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "description": "Brief description",
            "technologies": ["Tech1", "Tech2"],
            "highlights": ["outcome 1"]
        }}
    ],
    "languages": ["English", "Spanish"],
    "certifications": ["Cert1", "Cert2"]
}}"""

        try:
            parsed = await self._call_claude(system_prompt, user_prompt, max_tokens=3000)

            return ParsedResume(
                name=parsed.get("name"),
                email=parsed.get("email"),
                phone=parsed.get("phone"),
                location=parsed.get("location"),
                summary=parsed.get("summary"),
                skills=parsed.get("skills", []),
                experience=[
                    ExperienceItem(**exp) for exp in parsed.get("experience", [])
                ],
                education=[
                    EducationItem(**edu) for edu in parsed.get("education", [])
                ],
                projects=[
                    ProjectItem(**proj) for proj in parsed.get("projects", [])
                ],
                languages=parsed.get("languages", []),
                certifications=parsed.get("certifications", [])
            )

        except Exception as e:
            logger.error(f"Claude resume parsing error: {e}")
            # Try DeepSeek fallback
            if self.deepseek_api_key:
                return await self._parse_resume_deepseek(raw_text)
            return ParsedResume()

    async def _parse_resume_deepseek(self, raw_text: str) -> ParsedResume:
        """Fallback to DeepSeek for resume parsing."""
        if not self.deepseek_api_key:
            return ParsedResume()

        system_prompt = """You are an expert resume parser. Extract structured information accurately.
Respond in valid JSON format only."""

        user_prompt = f"""Parse this resume:

{raw_text[:8000]}

Respond with JSON:
{{
    "name": "Full Name or null",
    "email": "email or null",
    "phone": "phone or null",
    "location": "location or null",
    "summary": "summary or null",
    "skills": ["skill1", "skill2"],
    "experience": [{{"company": "...", "title": "...", "start_date": "...", "end_date": "...", "description": "...", "highlights": []}}],
    "education": [{{"institution": "...", "degree": "...", "field_of_study": "...", "start_date": "...", "end_date": "...", "gpa": "..."}}],
    "projects": [{{"name": "...", "description": "...", "technologies": [], "highlights": []}}],
    "languages": [],
    "certifications": []
}}"""

        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{self.deepseek_base_url}/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.deepseek_api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "deepseek-chat",
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ],
                        "temperature": 0.1,
                        "max_tokens": 3000,
                        "response_format": {"type": "json_object"}
                    }
                )
                response.raise_for_status()
                result = response.json()
                content = result["choices"][0]["message"]["content"]
                parsed = json.loads(content)

                return ParsedResume(
                    name=parsed.get("name"),
                    email=parsed.get("email"),
                    phone=parsed.get("phone"),
                    location=parsed.get("location"),
                    summary=parsed.get("summary"),
                    skills=parsed.get("skills", []),
                    experience=[
                        ExperienceItem(**exp) for exp in parsed.get("experience", [])
                    ],
                    education=[
                        EducationItem(**edu) for edu in parsed.get("education", [])
                    ],
                    projects=[
                        ProjectItem(**proj) for proj in parsed.get("projects", [])
                    ],
                    languages=parsed.get("languages", []),
                    certifications=parsed.get("certifications", [])
                )

        except Exception as e:
            logger.error(f"DeepSeek resume parsing error: {e}")
            return ParsedResume()

    async def generate_personalized_questions(
        self,
        parsed_resume: ParsedResume,
        job_title: Optional[str] = None,
        job_requirements: Optional[list[str]] = None,
        num_questions: int = 3
    ) -> list[PersonalizedQuestion]:
        """
        Generate personalized interview questions based on resume content.
        Uses Claude for fast, high-quality question generation.

        Args:
            parsed_resume: Structured resume data
            job_title: Target job title
            job_requirements: List of job requirements
            num_questions: Number of questions to generate

        Returns:
            List of PersonalizedQuestion objects
        """
        if not self.anthropic_api_key and not self.deepseek_api_key:
            return []

        # Build context from resume
        resume_context = []
        if parsed_resume.experience:
            recent_exp = parsed_resume.experience[0]
            resume_context.append(f"Recent role: {recent_exp.title} at {recent_exp.company}")
        if parsed_resume.skills:
            resume_context.append(f"Skills: {', '.join(parsed_resume.skills[:10])}")
        if parsed_resume.projects:
            resume_context.append(f"Notable project: {parsed_resume.projects[0].name}")
        if parsed_resume.education:
            edu = parsed_resume.education[0]
            resume_context.append(f"Education: {edu.degree} from {edu.institution}")

        job_context = job_title or "General"
        requirements_text = "\n".join(f"- {req}" for req in (job_requirements or []))

        system_prompt = """You are an expert interviewer. Generate insightful, personalized interview questions.
Questions should probe deeper into the candidate's experience and assess role fit.
Respond with valid JSON only."""

        user_prompt = f"""Generate {num_questions} personalized interview questions:

CANDIDATE:
{chr(10).join(resume_context)}

TARGET: {job_context}
{f"REQUIREMENTS:{chr(10)}{requirements_text}" if requirements_text else ""}

Return JSON:
{{
    "questions": [
        {{
            "text": "English question",
            "text_zh": "Chinese translation",
            "category": "behavioral/technical/experience",
            "based_on": "Resume element referenced"
        }}
    ]
}}"""

        try:
            if self.anthropic_api_key:
                parsed = await self._call_claude(system_prompt, user_prompt, max_tokens=2000)
            else:
                # Fallback to DeepSeek
                async with httpx.AsyncClient(timeout=60.0) as client:
                    response = await client.post(
                        f"{self.deepseek_base_url}/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {self.deepseek_api_key}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": "deepseek-chat",
                            "messages": [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_prompt}
                            ],
                            "temperature": 0.7,
                            "max_tokens": 2000,
                            "response_format": {"type": "json_object"}
                        }
                    )
                    response.raise_for_status()
                    result = response.json()
                    content = result["choices"][0]["message"]["content"]
                    parsed = json.loads(content)

            return [
                PersonalizedQuestion(**q)
                for q in parsed.get("questions", [])
            ]

        except Exception as e:
            logger.error(f"Question generation error: {e}")
            return []

    async def generate_adaptive_questions(
        self,
        candidate_context: dict,
        parsed_resume: Optional[ParsedResume] = None,
        job_title: Optional[str] = None,
        job_requirements: Optional[list[str]] = None,
        num_questions: int = 3
    ) -> list[PersonalizedQuestion]:
        """
        Generate ADAPTIVE interview questions using ALL available candidate data.
        This creates truly personalized questions based on:
        - Resume (experience, projects, skills)
        - GitHub (repos, languages, contributions)
        - Transcript (courses, grades, academic trajectory)

        Args:
            candidate_context: Dict with candidate data from all sources
            parsed_resume: Optional parsed resume data
            job_title: Target job title
            job_requirements: List of job requirements
            num_questions: Number of questions to generate

        Returns:
            List of PersonalizedQuestion objects
        """
        if not self.anthropic_api_key and not self.deepseek_api_key:
            # Fall back to basic personalized questions if we have a resume
            if parsed_resume:
                return await self.generate_personalized_questions(
                    parsed_resume, job_title, job_requirements, num_questions
                )
            return []

        # Build comprehensive context string
        context_parts = []

        # Basic info
        if candidate_context.get("name"):
            context_parts.append(f"Name: {candidate_context['name']}")
        if candidate_context.get("university"):
            context_parts.append(f"University: {candidate_context['university']}")
        if candidate_context.get("majors"):
            majors = candidate_context["majors"]
            if len(majors) > 1:
                context_parts.append(f"Double Major: {' & '.join(majors)}")
            elif majors:
                context_parts.append(f"Major: {majors[0]}")
        elif candidate_context.get("major"):
            context_parts.append(f"Major: {candidate_context['major']}")
        if candidate_context.get("graduation_year"):
            context_parts.append(f"Graduation: {candidate_context['graduation_year']}")
        if candidate_context.get("gpa"):
            context_parts.append(f"GPA: {candidate_context['gpa']}")

        # Resume data
        if candidate_context.get("resume"):
            resume = candidate_context["resume"]
            if resume.get("skills"):
                context_parts.append(f"\nResume Skills: {', '.join(resume['skills'][:10])}")
            if resume.get("experiences"):
                exp_strs = []
                for exp in resume["experiences"][:2]:
                    exp_str = f"{exp.get('title', 'Role')} at {exp.get('company', 'Company')}"
                    if exp.get("highlights"):
                        exp_str += f" - {exp['highlights'][0]}"
                    exp_strs.append(exp_str)
                context_parts.append(f"Experience: {'; '.join(exp_strs)}")
            if resume.get("projects"):
                proj_strs = []
                for proj in resume["projects"][:2]:
                    proj_str = f"{proj.get('name', 'Project')}"
                    if proj.get("technologies"):
                        proj_str += f" ({', '.join(proj['technologies'][:3])})"
                    proj_strs.append(proj_str)
                context_parts.append(f"Projects: {'; '.join(proj_strs)}")

        # GitHub data (NEW - enables coding-specific questions)
        if candidate_context.get("github"):
            github = candidate_context["github"]
            context_parts.append(f"\n--- GitHub Profile ---")
            if github.get("username"):
                context_parts.append(f"GitHub: @{github['username']}")
            if github.get("top_languages"):
                context_parts.append(f"Top Languages: {', '.join(github['top_languages'])}")
            if github.get("top_repos"):
                repo_strs = []
                for repo in github["top_repos"]:
                    repo_str = f"{repo.get('name', 'repo')}"
                    if repo.get("language"):
                        repo_str += f" ({repo['language']})"
                    if repo.get("stars", 0) > 0:
                        repo_str += f" ⭐{repo['stars']}"
                    repo_strs.append(repo_str)
                context_parts.append(f"Notable Repos: {', '.join(repo_strs)}")
            if github.get("contributions"):
                context_parts.append(f"Total Contributions: {github['contributions']}")

        # Transcript data (NEW - enables course-specific questions)
        if candidate_context.get("transcript"):
            transcript = candidate_context["transcript"]
            context_parts.append(f"\n--- Academic Record ---")
            if transcript.get("gpa"):
                context_parts.append(f"Cumulative GPA: {transcript['gpa']}")
            if transcript.get("technical_gpa"):
                context_parts.append(f"Technical GPA: {transcript['technical_gpa']}")
            if transcript.get("top_courses"):
                context_parts.append(f"Strong Courses (A/A+): {', '.join(transcript['top_courses'])}")
            if transcript.get("transcript_score"):
                context_parts.append(f"Academic Score: {transcript['transcript_score']}/100")
        elif candidate_context.get("courses"):
            context_parts.append(f"Relevant Courses: {', '.join(candidate_context['courses'][:8])}")

        full_context = "\n".join(context_parts)
        job_context = job_title or "Software Engineering"
        requirements_text = "\n".join(f"- {req}" for req in (job_requirements or []))

        system_prompt = """You are an expert technical interviewer for US tech companies.
Generate highly personalized interview questions that:
1. Reference SPECIFIC items from the candidate's profile (a GitHub repo, a course, a project)
2. Assess both technical depth and problem-solving ability
3. Are appropriate for a college student/new grad level
4. Mix behavioral and technical questions

If GitHub data is available, ask about their actual repositories/contributions.
If transcript data is available, ask about challenging courses they excelled in.
If project data is available, dig deeper into their implementations.

Respond with valid JSON only."""

        user_prompt = f"""Generate {num_questions} highly personalized interview questions for this candidate:

=== CANDIDATE PROFILE ===
{full_context}

=== TARGET ROLE ===
{job_context}
{f"Requirements:{chr(10)}{requirements_text}" if requirements_text else ""}

=== INSTRUCTIONS ===
Create questions that SPECIFICALLY reference their profile:
- If they have a notable GitHub repo, ask about its architecture or a technical decision
- If they took a hard course (like CS 61B or OS), ask how they'd apply that knowledge
- If they have project experience, ask about scaling it or handling edge cases
- If they're a double major, ask how their diverse background helps them

Return JSON:
{{
    "questions": [
        {{
            "text": "English question that references specific profile item",
            "text_zh": "Chinese translation",
            "category": "technical/behavioral/experience",
            "based_on": "The specific profile element this question references"
        }}
    ]
}}"""

        try:
            if self.anthropic_api_key:
                parsed = await self._call_claude(system_prompt, user_prompt, max_tokens=2000)
            else:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    response = await client.post(
                        f"{self.deepseek_base_url}/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {self.deepseek_api_key}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": "deepseek-chat",
                            "messages": [
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_prompt}
                            ],
                            "temperature": 0.7,
                            "max_tokens": 2000,
                            "response_format": {"type": "json_object"}
                        }
                    )
                    response.raise_for_status()
                    result = response.json()
                    content = result["choices"][0]["message"]["content"]
                    parsed = json.loads(content)

            questions = [
                PersonalizedQuestion(**q)
                for q in parsed.get("questions", [])
            ]

            logger.info(f"Generated {len(questions)} adaptive questions based on: "
                       f"resume={bool(candidate_context.get('resume'))}, "
                       f"github={bool(candidate_context.get('github'))}, "
                       f"transcript={bool(candidate_context.get('transcript'))}")

            return questions

        except Exception as e:
            logger.error(f"Adaptive question generation error: {e}")
            # Fall back to basic personalized questions
            if parsed_resume:
                return await self.generate_personalized_questions(
                    parsed_resume, job_title, job_requirements, num_questions
                )
            return []


# Global instance
resume_service = ResumeService()
